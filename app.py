from flask import Flask, jsonify, render_template, request, send_file
import os
import json
import os
import re
from docx import Document
from markdownify import markdownify as md
from datetime import datetime
from openai_utils import gpt_4o_openai_functon, read_markdown_file, save_to_markdown, text_mode_message_list_example
from flask_cors import CORS
from build_docker import package_flask_app  

app = Flask(__name__, static_folder='fronted/')
CORS(app)
current_dir = "./my_data"
time_out_of_date = 20

@app.route('/')
def index():
    return render_template('index.html')


def step_1_convert_docx_to_md(docx_path, md_dir):
    # Load the .docx file
    doc = Document(docx_path)
   
    # Read all paragraphs and combine them into a single string
    doc_text = "\n".join([p.text for p in doc.paragraphs])
   
    # Convert the text to markdown
    markdown_text = md(doc_text)
   
    # Write the markdown text to a .md file
    os.makedirs(md_dir, exist_ok=True)
    md_path = os.path.join(md_dir, "1.md")
    with open(md_path, 'w') as md_file:
        md_file.write(markdown_text)


def step_2_analyse(md_dir):
    a = """
    You are a highly efficient assistant who will extract and analyze the information from the provided markdown file.
    Please pay close attention to details and the logical structure within the markdown file.
    Your task is to group related information together and make the content more insightful.

    Design survey questions that are unbiased and easy to collect data from without drawing attention.
    Plan ahead to anticipate expected outputs and the intentions behind them.
    Ensure that the output is structured and modular, and format it in markdown for easy review.

    The markdown output should include Well-crafted, unbiased survey questions.
    """

    os.makedirs(md_dir, exist_ok=True)
    last_step_md_path = os.path.join(md_dir, "1.md")
    b = f'{read_markdown_file(last_step_md_path)}'
    result = gpt_4o_openai_functon(text_mode_message_list_example(a, b))
    result_path =  os.path.join(md_dir, "2.md")
    save_to_markdown(result, result_path)


def step_3_get_survey_md(md_dir):
    second_step_md_file =  os.path.join(md_dir, "2.md")
    content = read_markdown_file(second_step_md_file)
    a = f"""
You are a highly experienced and helpful assistant with extensive expertise in designing survey questions.
Your goal is to minimize bias and provide options that achieve this.
Example questions are provided for you to reference the structures.
For each question, please include the intention, layer, and expected output,
so it is easy for me to review.
Ensure the output is in markdown format.

The output should include:
1. Well-structured survey questions with options.
2. Each option needs to be separated by a hyphen and a newline.
3. The intention behind each question.
4. The layer each question belongs to.
5. The expected output from each question.
6. The question type from ['Close-ended: Single option', 'Close-ended: Multiple options', 'Open-ended: Text answer']

These questions should be new and not taken directly from existing samples, but should have similar intentions and structure to effectively complete the survey. Each question should be well-constructed to minimize bias and provide clear, actionable insights. Additionally, please specify the question type as one of the following:
- Close-ended: Single option
- Close-ended: Multiple options
- Open-ended: Text answer

and the EXAMPLE is like
   
{content}
   
"""

    b = """
    Please help me design 20 survey questions related to the topics of 綠色金融 (Green Finance) and 新質生產力 (New Quality Productivity).
    """

    result = gpt_4o_openai_functon(text_mode_message_list_example(a, b))
    result_path =  os.path.join(md_dir, "3.md")
    save_to_markdown(result, result_path)

def step_4_extract_structure(current_dir):
    PATH_TO_INPUT_MARKDOWN = os.path.join(current_dir, "3.md")
    PATH_TO_OUTPUT_FILE = os.path.join(current_dir, "4.json")
    # Regular expressions to match sections, subsections, and questions
    section_regex = re.compile(r"^# (.+)")
    subsection_regex = re.compile(r"^## (.+)")
    subsubsection_regex = re.compile(r"^### (.+)")
    question_regex = re.compile(r"^\d+\. (.+)")
   
    question_intention_regex = re.compile(r"^\*\*Intention:\*\* +(.*)(?:\s|$)")
    question_layer_regex = re.compile(r"^\*\*Layer:\*\* +(.*)(?:\s|$)")
    question_output_regex = re.compile(r"^\*\*Expected Output:\*\* +(.*)(?:\s|$)")
    question_type_regex = re.compile(r"^\*\*Question Type:\*\* +(.*)(?:\s|$)")
   
    # Place holders
    structure = []
    current_section = None
    current_subsection = None
    current_subsubsection = None
   
    with open(PATH_TO_INPUT_MARKDOWN, 'r', encoding='utf-8') as file:
        lines = file.readlines()
        current_question = None
        for line in lines:
            line = line.strip()
           
            if m := section_regex.match(line):
                current_section = {'section': m.group(1), 'subsections': []}
                structure.append(current_section)
                current_subsection = None
                current_subsubsection = None
            elif m := subsection_regex.match(line):
                current_subsection = {'subsection': m.group(1), 'questions': []}
                if current_section:
                    current_section['subsections'].append(current_subsection)
                current_subsubsection = None
            elif m := subsubsection_regex.match(line):
                current_subsubsection = {'subsubsection': m.group(1), 'questions': []}
                if current_subsection:
                    current_subsection['questions'].append(current_subsubsection)
                elif current_section:
                    current_section['subsections'].append(current_subsubsection)
            elif m := question_regex.match(line):
                current_question = {'question': m.group(1), 'options': [], 'intention': None, 'layer': None, 'expected_output': None, 'question_type': None}
                if current_subsubsection:
                    current_subsubsection['questions'].append(current_question)
                elif current_subsection:
                    current_subsection['questions'].append(current_question)
                elif current_section:
                    current_section['subsections'].append(current_question)
            elif re.match(r"^-", line):
                if current_question:
                    current_question['options'].append(line.lstrip('- '))
            elif m := question_intention_regex.match(line):
                if current_question:
                    current_question['intention'] = m.group(1)
            elif m := question_layer_regex.match(line):
                if current_question:
                    current_question['layer'] = m.group(1)
            elif m := question_output_regex.match(line):
                if current_question:
                    current_question['expected_output'] = m.group(1)
            elif m := question_type_regex.match(line):
                if current_question:
                    current_question['question_type'] = m.group(1)

    with open(PATH_TO_OUTPUT_FILE, 'w', encoding='utf-8') as outfile:
        json.dump(structure, outfile, ensure_ascii=False, indent=4)


def extract_python_code_from_markdown(current_dir):
    input_markdown_path = os.path.join(current_dir, "4.md")
    work_dir = os.path.join(current_dir, "workspace")
    os.makedirs(work_dir, exist_ok=True)
    output_python_path  = os.path.join(work_dir, "app.py")
    # Regular expression to find code blocks
    code_block_regex = re.compile(r"```python(.*?)```", re.DOTALL)
   
    with open(input_markdown_path, 'r', encoding='utf-8') as file:
        markdown_content = file.read()
   
    # Find all Python code blocks
    python_code_blocks = code_block_regex.findall(markdown_content)
   
    # Write the extracted code blocks to the output file
    with open(output_python_path, 'w', encoding='utf-8') as output_file:
        for block in python_code_blocks:
            output_file.write(block.strip() + '\n\n')


def step_4_get_python_extract_code(current_dir):
    PATH_TO_INPUT_MARKDOWN = os.path.join(current_dir, "3.md")
    PATH_TO_OUTPUT_FILE = os.path.join(current_dir, "./4.json")
    a = """
I need a Python function to extract the structure from a markdown (.md) file. The function should take two variables:
1. `PATH_TO_INPUT_MARKDOWN` - the path to the input markdown file and the default value is {PATH_TO_INPUT_MARKDOWN}
2. `PATH_TO_OUTPUT_FILE` - the path to the output file and the default value is {PATH_TO_OUTPUT_FILE}

The function should read the markdown file, extract the structure including section, subsection, question(including the options, type , number, Intention, Layer, Expected Output of the question)
the questions belong to the subsections, the subsections belong to the section,
then save the extracted structure to the specified output file in JSON format.

the input markdown is like the following user input

"""
    md_path = os.path.join(current_dir, "3.md")
    b = read_markdown_file(md_path)
    result = gpt_4o_openai_functon(text_mode_message_list_example(a, b))
    result_path =  os.path.join(current_dir, "4.md")
    save_to_markdown(result, result_path)


def get_flask_code(current_dir):
    a = """
    我需要你帮我我创建一个调查问卷的flask程序，用到render_template_string，样式要现代且精美，加一些背景动画，要有滚动条，前后端代码都在一个文件中, 程序运行的ip是'0.0.0.0'，端口是5000端口上。
    问卷数据的模板格式是markdown,我会在下面的用户输入给到你问卷数据具体的内容。
    你需要根据我给你的问卷数据的格式和内容设计问卷。
    """
    md_path = os.path.join(current_dir, "3.md")
    b = read_markdown_file(md_path)
    result = gpt_4o_openai_functon(text_mode_message_list_example(a, b))
    result_path =  os.path.join(current_dir, "4.md")
    save_to_markdown(result, result_path)

@app.route('/generate_preview', methods=['POST'])
def generate_preview():
    os.makedirs(current_dir, exist_ok=True)
    # 获取请求中的JSON数据
    data = request.json
    content = data.get('content')
    word_file_path = os.path.join(current_dir, "init.docx")
    if not content or not word_file_path:
        return jsonify({'error': 'markdown_content and word_file_path are required'}), 400

    # 创建一个新的Word文档
    doc = Document()

    # 将Markdown内容逐行添加到Word文档中
    for line in content.split('\n'):
        doc.add_paragraph(line)

    # 保存Word文档到指定路径
    doc.save(word_file_path)
    step_1_convert_docx_to_md(word_file_path, current_dir)
    step_2_analyse(current_dir)
    step_3_get_survey_md(current_dir)
    preview_file_path = os.path.join(current_dir, "3.md")
    preview_content = read_markdown_file(preview_file_path)
    return jsonify({"preview_content": preview_content, "current_dir": current_dir})

@app.route('/deploy', methods=['POST'])
def deploy():
    get_flask_code(current_dir)
    extract_python_code_from_markdown(current_dir)
    workspace_dir = os.path.join(current_dir, "workspace")
    address = package_flask_app(workspace_dir, "test")
    address = "https://supreme-zebra-w9qp4xrgvjgfg6v7-9987.app.github.dev/"
    return jsonify({"address": address})

if __name__ == '__main__':
    app.run(host="0.0.0.0", port="8889")